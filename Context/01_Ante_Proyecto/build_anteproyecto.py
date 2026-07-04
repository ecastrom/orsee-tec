# -*- coding: utf-8 -*-
"""
Genera la Propuesta de Construcción de Know-how y Plan de Uso del BEER Lab.
Salida: Ante_Proyecto_BEER_Lab_Docencia.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\L03574200\Dropbox\BEER - LAB\01_Ante_Proyecto\Ante_Proyecto_BEER_Lab_Docencia.docx"

doc = Document()

# ---------- Estilos base ----------
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for level, size in [(1, 16), (2, 13), (3, 12)]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ---------- Helpers ----------
def add_para(text, bold=False, italic=False, align=None, size=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level:
        p.paragraph_format.left_indent = Cm(0.6 * (level + 1))
    return p


def add_numbered(text):
    return doc.add_paragraph(text, style='List Number')


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_table(headers, rows, col_widths_cm=None, header_color="1F3A68"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = False

    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_text(cell, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, header_color)

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], str(val), size=10)
    doc.add_paragraph()
    return table


# ====================================================================
# PORTADA
# ====================================================================
top = doc.add_paragraph()
top.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = top.add_run("PROPUESTA OPERATIVA")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Construcción de Know-how y Plan de Uso\ndel BEER Lab")
r.bold = True
r.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Brazo Docente  |  Behavioral and Experimental Economics Research Lab")
r.font.size = Pt(11)
r.italic = True

org = doc.add_paragraph()
org.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = org.add_run("Departamento de Economía  |  Tecnológico de Monterrey, Campus Monterrey")
r.font.size = Pt(11)
r.italic = True

doc.add_paragraph()

meta = doc.add_table(rows=6, cols=2)
meta.style = 'Light List Accent 1'
meta_rows = [
    ("Responsables", "Edgar Castro Méndez (ecastrom@tec.mx)\nStanislao Maldonado Zambrano (stanislao.maldonado@tec.mx)"),
    ("Modelo educativo", "Tec21 — Aprendizaje vivencial aplicado a Unidades de Formación de Economía"),
    ("Recursos institucionales comprometidos", "Adecuación física del laboratorio, equipamiento de cómputo y audiovisual, servidor ORSEE, y laboratorio móvil (cubiertos por la institución)"),
    ("Horizonte del proyecto", "Tres semestres consecutivos: Ago–Dic 2026, Feb–Jun 2027, Ago–Dic 2027"),
    ("Equipo operativo", "10 a 12 estudiantes de Servicio Becario por semestre (80 horas por plaza), activos durante el semestre académico"),
    ("Documento", "Propuesta operativa — mayo 2026"),
]
for i, (k, v) in enumerate(meta_rows):
    set_cell_text(meta.rows[i].cells[0], k, bold=True, size=10)
    set_cell_text(meta.rows[i].cells[1], v, size=10)
    shade_cell(meta.rows[i].cells[0], "E8EEF7")

doc.add_paragraph()

# ====================================================================
# RESUMEN EJECUTIVO
# ====================================================================
add_heading("Resumen ejecutivo", level=1)
add_para(
    "La institución ha aprobado la creación del BEER Lab y financiará tanto la adecuación de la sala "
    "como el equipamiento del laboratorio fijo y del laboratorio móvil. Este documento parte de ese "
    "compromiso institucional y se ocupa de lo que sigue: construir el know-how y dejar instalada la "
    "capacidad operativa necesaria para que el laboratorio cumpla su función pedagógica desde el primer "
    "día que un docente del Departamento solicite una sesión. La infraestructura física es necesaria pero "
    "no suficiente; el valor educativo depende de tener un portafolio curado de experimentos en oTree, "
    "un banco de participantes funcional, un equipo formado, materiales de difusión, y reglas de uso claras.",
    space_after=8,
)
add_para(
    "La propuesta se ejecuta en tres semestres consecutivos (Ago–Dic 2026, Feb–Jun 2027, Ago–Dic 2027), "
    "con un equipo de 10 a 12 estudiantes de Servicio Becario por semestre organizados en cinco roles. "
    "Cada plaza cubre 80 horas de servicio durante el semestre, conforme a los lineamientos del programa "
    "institucional. El primer semestre construye los componentes base; el segundo pilotea, lanza el "
    "servicio y opera las primeras sesiones a solicitud de docentes; el tercero estabiliza la operación. "
    "Entre semestres —periodo invernal y verano— los Servicios Becarios no están activos; esos huecos "
    "se aprovechan para revisión, documentación y planeación por parte de los profesores responsables.",
    space_after=8,
)
add_para(
    "Al cierre del horizonte planteado, el Departamento de Economía contará con: un portafolio de al menos "
    "doce experimentos clásicos en oTree listos para integrarse a Unidades de Formación; una matriz "
    "curricular que mapea cada UF con experimentos pertinentes y competencias Tec21 activadas; un banco "
    "de participantes operativo; identidad visual y materiales de difusión; un manual de operación; y "
    "una rotación sostenida de estudiantes formados en metodología experimental, programación oTree, "
    "gestión de laboratorio y comunicación académica.",
    space_after=8,
)

# ====================================================================
# 1. MARCO INSTITUCIONAL
# ====================================================================
add_heading("1. Marco institucional", level=1)
add_para(
    "Esta propuesta se enmarca en la decisión institucional ya tomada de crear el BEER Lab como "
    "infraestructura docente del Departamento de Economía del Campus Monterrey. La justificación "
    "pedagógica, el posicionamiento estratégico, el diseño físico de la sala, el presupuesto de equipamiento "
    "y la inclusión del laboratorio móvil están documentados en la propuesta original (Castro y Maldonado, "
    "2026) y en el canvas de modelo de negocio social del laboratorio. No se reabren aquí esos puntos.",
    space_after=6,
)
add_para(
    "Lo que este documento aporta es la pieza operativa complementaria: el plan para construir el "
    "know-how que la infraestructura no trae consigo, y la propuesta de uso que define cómo el "
    "laboratorio se ofrecerá al claustro del Departamento como servicio integrado al modelo Tec21.",
    space_after=6,
)

# ====================================================================
# 2. OBJETO Y ALCANCE
# ====================================================================
add_heading("2. Objeto y alcance del documento", level=1)
add_para("Este documento define:", space_after=2)
add_bullet("Los componentes de capacidad operativa (know-how) que es necesario construir para que el BEER Lab cumpla su función pedagógica.")
add_bullet("La propuesta de uso del laboratorio: quién lo puede usar, bajo qué modalidades, con qué prioridades, y a través de qué proceso.")
add_bullet("Las fases del proyecto, alineadas al calendario académico del Campus Monterrey, reconociendo que los Servicios Becarios solo están activos durante el semestre.")
add_bullet("La composición del equipo de Servicio Becario por cohorte semestral, con cinco roles diferenciados y mecanismos de continuidad entre semestres.")
add_bullet("Los indicadores de éxito, la gobernanza interna y los riesgos relevantes para la ejecución.")
doc.add_paragraph()
add_para(
    "Fuera de alcance: el detalle de la inversión institucional en infraestructura (cubierto en la "
    "propuesta original); el diseño del brazo de investigación del BEER Lab (documento aparte); y la "
    "definición de productos de vinculación con la industria vía Expedition (mediano plazo).",
    space_after=6,
)

# ====================================================================
# 3. COMPONENTES DEL KNOW-HOW A CONSTRUIR
# ====================================================================
add_heading("3. Componentes del know-how a construir", level=1)
add_para(
    "El valor pedagógico del laboratorio depende de cinco componentes que la infraestructura física no "
    "trae consigo y que esta propuesta se compromete a entregar.",
    space_after=8,
)

add_heading("3.1. Portafolio curado de experimentos en oTree", level=2)
add_para(
    "El primer componente es un catálogo de experimentos clásicos transcritos a oTree, con parámetros "
    "calibrados, instrucciones que no priman al participante, predicciones teóricas documentadas, y "
    "fichas técnicas didácticas que permitan a cualquier profesor titular elegir el experimento adecuado "
    "para su Unidad de Formación. Meta del horizonte: doce experimentos completos, piloteados y listos "
    "para producción.",
    space_after=6,
)

add_heading("3.2. Mapeo curricular UF × experimento × competencia Tec21", level=2)
add_para(
    "Una matriz que asocia cada experimento del portafolio con las Unidades de Formación del Departamento "
    "en las que es pertinente y con las competencias —disciplinares y transversales del Modelo Tec21— "
    "que la sesión experimental activa. Sin este mapeo, el catálogo es solo una lista; con él, se vuelve "
    "una herramienta de planeación pedagógica que los profesores titulares pueden integrar a sus "
    "planeaciones semestrales.",
    space_after=6,
)

add_heading("3.3. Banco de participantes operativo (ORSEE)", level=2)
add_para(
    "Una base de datos institucional de estudiantes registrados como potenciales participantes, "
    "soportada por la plataforma de código abierto ORSEE (Online Recruitment System for Economic "
    "Experiments). Permite convocatorias automatizadas, registro de participación histórica, y "
    "selección de muestras por criterios demográficos. Sin este componente, cada sesión exige "
    "reclutamiento manual, lo que vuelve insostenible el servicio.",
    space_after=6,
)

add_heading("3.4. Identidad visual y kit de difusión", level=2)
add_para(
    "El servicio necesita ser visible y reconocible dentro del Departamento. Esto implica una identidad "
    "visual básica (logo, paleta, plantillas), un flyer que comunique la oferta en menos de un minuto "
    "de lectura, una página interna informativa, y una sesión informativa para profesores titulares "
    "en el momento del lanzamiento.",
    space_after=6,
)

add_heading("3.5. Manual de operación", level=2)
add_para(
    "El protocolo escrito que estandariza la operación: cómo se solicita una sesión, cómo se prepara "
    "la sala, cómo se ejecuta el experimento, cómo se procesan los pagos a participantes, qué reglas de "
    "consentimiento informado y protección de datos aplican. El manual garantiza que la calidad del "
    "servicio no dependa de la memoria de un equipo en particular y soporta la rotación de Servicios "
    "Becarios entre semestres.",
    space_after=6,
)

# ====================================================================
# 4. PROPUESTA DE USO
# ====================================================================
add_heading("4. Propuesta de uso del laboratorio", level=1)
add_para(
    "El BEER Lab opera como servicio del Departamento de Economía. La propuesta de uso define "
    "beneficiarios, modalidades, proceso de solicitud, prioridades y reglas.",
    space_after=8,
)

add_heading("4.1. Beneficiarios", level=2)
add_para("Beneficiarios primarios (primer ciclo):", bold=True, space_after=2)
add_bullet("Profesores titulares del Departamento de Economía con Unidades de Formación susceptibles de integrar sesiones experimentales (ver sección 5).")
add_bullet("Estudiantes inscritos en esas Unidades de Formación, en calidad de participantes en las sesiones experimentales integradas a su clase.")
add_para("Beneficiarios secundarios (a partir del ciclo 2):", bold=True, space_after=2)
add_bullet("Profesores titulares de otros departamentos (Negocios, Mercadotecnia, Políticas Públicas, Ciencias Sociales) interesados en aprendizaje vivencial vía experimentos económicos.")
add_bullet("Bloques integradores y bloques multidisciplinares en los que la economía conductual aporta un componente experimental.")
add_bullet("Investigadores internos del Campus con proyectos de investigación experimental aprobados.")
add_bullet("Estudiantes que participen como Servicios Becarios y reciban formación en metodología experimental.")
doc.add_paragraph()

add_heading("4.2. Modalidades de uso", level=2)
add_para("Sesión integrada a UF.", bold=True, space_after=2)
add_para(
    "Modalidad estándar. Un profesor titular solicita una sesión de aproximadamente 90 minutos dentro "
    "de su clase regular, durante la cual sus estudiantes participan como sujetos en un experimento "
    "seleccionado del portafolio. El experimento se acompaña de una discusión post-sesión guiada por el "
    "profesor titular, en la que se conectan los resultados observados con el contenido teórico de la UF.",
    space_after=6,
)
add_para("Bloque experimental.", bold=True, space_after=2)
add_para(
    "Para Bloques integradores o Bloques multidisciplinares que requieran un componente experimental "
    "más profundo, se ofrece una secuencia de varias sesiones complementarias a lo largo del periodo, "
    "diseñadas en conjunto con el profesor responsable del bloque.",
    space_after=6,
)
add_para("Sesión móvil.", bold=True, space_after=2)
add_para(
    "Cuando el laboratorio principal no esté disponible o el grupo sea demasiado grande, el laboratorio "
    "móvil se traslada al aula correspondiente. Esta modalidad amplía el alcance del servicio sin "
    "saturar la sala fija de Aulas VI.",
    space_after=6,
)
add_para("Sesión de investigación.", bold=True, space_after=2)
add_para(
    "A partir del segundo ciclo, fuera del horario de uso prioritario docente, el laboratorio puede "
    "atender solicitudes de investigadores internos o externos con protocolos aprobados por el comité "
    "ético correspondiente.",
    space_after=6,
)

add_heading("4.3. Proceso de solicitud", level=2)
add_numbered("El profesor titular llena un formulario en línea con la UF, el número aproximado de estudiantes, el experimento de interés del portafolio (o solicita una recomendación) y dos o tres fechas tentativas.")
add_numbered("El Lab Manager verifica disponibilidad de sala, equipo y participantes (si requiere reclutamiento adicional vía ORSEE), y propone una agenda confirmada.")
add_numbered("La sesión queda agendada con al menos tres semanas de anticipación; en ese plazo el equipo del laboratorio prepara materiales, validaciones técnicas y, si aplica, la convocatoria a participantes externos.")
add_numbered("El día de la sesión, el equipo de Servicios Becarios opera la sesión bajo el protocolo del manual. El profesor titular acompaña la sesión y conduce la discusión post-experimento.")
add_numbered("Tras la sesión, el laboratorio entrega al profesor titular un breve reporte de datos agregados (sin información personal identificable) que puede utilizar para la discusión en clase o para los productos pedagógicos posteriores.")
doc.add_paragraph()

add_heading("4.4. Prioridades de uso durante el primer ciclo", level=2)
add_bullet("Prioridad alta: Unidades de Formación impartidas por profesores titulares del Departamento de Economía.")
add_bullet("Prioridad media: Bloques integradores y multidisciplinares con componente económico explícito.")
add_bullet("Prioridad baja durante el primer año: otros departamentos e investigación externa. Se atienden conforme a disponibilidad y bajo acuerdo específico.")
doc.add_paragraph()

add_heading("4.5. Reglas de uso", level=2)
add_bullet("Reserva con al menos tres semanas de anticipación.")
add_bullet("Cada Unidad de Formación puede solicitar hasta dos sesiones por semestre durante el primer ciclo (regla ajustable conforme se acumule capacidad operativa).")
add_bullet("Los incentivos monetarios a participantes se cubren conforme al acuerdo administrativo vigente; no se realizan sesiones sin el mecanismo de pago resuelto.")
add_bullet("Toda sesión opera con consentimiento informado escrito de los participantes.")
add_bullet("El laboratorio no se utiliza como instrumento de evaluación individual; las sesiones son herramienta pedagógica, no de calificación.")
add_bullet("Los datos generados son propiedad institucional. Cualquier uso para publicación o difusión externa requiere acuerdo previo con los responsables del laboratorio y revisión ética cuando aplique.")
doc.add_paragraph()

# ====================================================================
# 5. CARTERA DE EXPERIMENTOS POR UF
# ====================================================================
add_heading("5. Cartera de experimentos por Unidad de Formación", level=1)
add_para(
    "El siguiente mapa preliminar asocia experimentos clásicos a Unidades de Formación del Departamento "
    "de Economía, con la competencia que el aprendizaje vivencial activa en cada caso. Es un punto de "
    "partida; la versión definitiva será validada por los analistas de mapeo curricular con base en los "
    "syllabi vigentes y entrevistas con profesores titulares durante el primer semestre del proyecto.",
    space_after=8,
)

uf_rows = [
    ("Economía conductual aplicada a las finanzas",
     "Bubble Market (Smith, Suchanek y Williams 1988); Asset trading con información asimétrica; Holt-Laury (aversión al riesgo); BDM elicitation.",
     "Razonamiento para la complejidad; pensamiento crítico; análisis de decisiones bajo incertidumbre."),
    ("Teoría de juegos y economía de la información",
     "Trust Game; Dictator Game; Ultimatum Game; Battle of the Sexes; Stag Hunt; Beauty Contest; Centipede Game.",
     "Razonamiento estratégico; competencia disciplinar en equilibrios y selección; comunicación de resultados."),
    ("Economía aplicada y herramientas computacionales",
     "Double Auction (Smith 1962); Bertrand y Cournot; subastas de primer y segundo precio; mecanismo de Vickrey.",
     "Inteligencia digital; razonamiento para la complejidad; competencia disciplinar en mecanismos de mercado."),
    ("Bloque integrador: Investigación aplicada — desafíos contemporáneos",
     "Replicación dirigida de un paper experimental; diseño guiado de un experimento propio; análisis y comunicación de hallazgos.",
     "Innovación y emprendimiento; comunicación; razonamiento para la complejidad; competencia disciplinar en metodología."),
    ("Bloque multidisciplinar (Economía + Mercadotecnia / Políticas Públicas)",
     "Nudges y arquitectura de decisiones; default effect; choice architecture; framing experiments.",
     "Conciencia y compromiso ciudadano; innovación; razonamiento para la complejidad."),
    ("Economía y finanzas personales aplicadas",
     "Time preferences (Convex Time Budget, Andreoni-Sprenger); risk preferences (Holt-Laury, Eckel-Grossman); mental accounting; status-quo bias.",
     "Autoconocimiento y gestión; vida saludable financiera; razonamiento para la complejidad."),
]
build_table(
    ["Unidad de Formación", "Experimentos del portafolio", "Competencias activadas (Tec21)"],
    uf_rows,
    col_widths_cm=[4.5, 6.0, 5.5],
)

add_para(
    "Cada experimento del portafolio se documenta con una ficha técnica didáctica que incluye: descripción "
    "teórica, parámetros recomendados, instrucciones para participantes, predicción teórica, resultados "
    "típicos en la literatura, sugerencias para la discusión post-experimento y duración estimada. La "
    "plantilla de ficha se construye en la primera mitad del Semestre Ago–Dic 2026, antes de iniciar la "
    "transcripción masiva al portafolio.",
    space_after=6,
)

# ====================================================================
# 6. FASES DEL PROYECTO
# ====================================================================
add_heading("6. Fases del proyecto", level=1)
add_para(
    "El proyecto se ejecuta en tres semestres consecutivos. Entre ellos, durante el periodo invernal "
    "(enero) y el verano (junio–agosto), los Servicios Becarios no están activos. Esos huecos se aprovechan "
    "para revisión, documentación, ajustes y planeación por parte de los profesores responsables.",
    space_after=8,
)

# Fase 0
add_heading("6.1. Fase 0 — Pre-arranque (Verano 2026, jun–ago)", level=2)
add_para("Sin Servicios Becarios activos.", italic=True, space_after=2)
add_para("Énfasis:", bold=True, space_after=2)
add_bullet("Coordinación con DTI, Dirección de Departamento y administración del Campus para calendarizar la adecuación física de la sala de Aulas VI y la entrega del laboratorio móvil.")
add_bullet("Convocatoria abierta de Servicio Becario para el semestre Ago–Dic 2026.")
add_bullet("Entrevistas y selección de la primera cohorte (cohorte 1).")
add_bullet("Diseño de la plantilla de ficha técnica didáctica y de la estructura del repositorio del portafolio.")
add_bullet("Preparación del onboarding de la cohorte 1.")
add_para("Entregables al cierre de fase:", bold=True, space_after=2)
add_bullet("Cohorte 1 conformada y notificada.")
add_bullet("Plantilla de ficha técnica didáctica v1.0.")
add_bullet("Estructura del repositorio del portafolio en git.")

# Fase 1
add_heading("6.2. Fase 1 — Construcción base (Semestre Ago–Dic 2026)", level=2)
add_para("Cohorte 1 de Servicios Becarios activa.", italic=True, space_after=2)
add_para("Énfasis:", bold=True, space_after=2)
add_bullet("Transcripción a oTree de los primeros cuatro a seis experimentos del portafolio.")
add_bullet("Instalación y configuración del servidor ORSEE; carga del primer padrón de potenciales participantes.")
add_bullet("Mapeo curricular completo: revisión de syllabi de todas las UFs del Departamento.")
add_bullet("Diseño y aplicación de la encuesta de interés docente; sistematización de resultados.")
add_bullet("Construcción del kit de identidad visual y borradores del flyer y la página interna.")
add_bullet("Definición del acuerdo administrativo con la Dirección Administrativa sobre pagos a participantes.")
add_para("Entregables al cierre del semestre:", bold=True, space_after=2)
add_bullet("Portafolio v1.0 con cuatro a seis experimentos funcionales y fichas técnicas validadas.")
add_bullet("Matriz UF × experimento × competencia Tec21 validada con el coordinador académico.")
add_bullet("Reporte ejecutivo de la encuesta de interés docente.")
add_bullet("Kit de identidad visual y flyer en borrador final.")
add_bullet("Página interna informativa publicada.")
add_bullet("Manual de operación v0.5 (borrador).")

# Receso invierno
add_heading("6.3. Receso invernal (enero 2027)", level=2)
add_para("Sin Servicios Becarios activos.", italic=True, space_after=2)
add_para("Énfasis (profesores responsables):", bold=True, space_after=2)
add_bullet("Revisión integral de los productos entregados por la cohorte 1.")
add_bullet("Ajustes finos al portafolio (fidelidad teórica, instrucciones, parámetros).")
add_bullet("Reunión de socialización con la Dirección del Departamento para presentar la matriz curricular y el reporte de encuesta.")
add_bullet("Planeación del lanzamiento del semestre Feb–Jun 2027.")
add_bullet("Convocatoria y selección de la cohorte 2, con retención prioritaria de las plazas ancla.")

# Fase 2
add_heading("6.4. Fase 2 — Pilotaje, lanzamiento y primeras sesiones (Semestre Feb–Jun 2027)", level=2)
add_para("Cohorte 2 de Servicios Becarios activa.", italic=True, space_after=2)
add_para("Énfasis:", bold=True, space_after=2)
add_bullet("Pilotaje del portafolio con grupos voluntarios para validación técnica y didáctica.")
add_bullet("Expansión del portafolio a doce experimentos.")
add_bullet("Sesión informativa formal para profesores titulares del Departamento.")
add_bullet("Apertura del calendario de solicitudes y operación de las primeras sesiones a solicitud docente.")
add_bullet("Documentación de los primeros casos de uso para difusión interna.")
add_bullet("Finalización del manual de operación.")
add_para("Entregables al cierre del semestre:", bold=True, space_after=2)
add_bullet("Portafolio v1.1 con doce experimentos validados.")
add_bullet("Manual de operación v1.0 publicado.")
add_bullet("Al menos cinco profesores titulares han solicitado y recibido una sesión experimental.")
add_bullet("Al menos quince sesiones experimentales ejecutadas (pilotajes más operativas).")
add_bullet("Casos de uso documentados para tres UFs.")

# Receso verano
add_heading("6.5. Receso de verano (jun–ago 2027)", level=2)
add_para("Sin Servicios Becarios activos.", italic=True, space_after=2)
add_para("Énfasis (profesores responsables):", bold=True, space_after=2)
add_bullet("Evaluación del primer año contra los indicadores de éxito.")
add_bullet("Reporte de cierre del primer año al Departamento y a la Decanatura.")
add_bullet("Ajustes al portafolio y al manual con base en aprendizajes operativos.")
add_bullet("Convocatoria y selección de la cohorte 3, con apertura limitada a otros departamentos.")
add_bullet("Planeación de la fase de estabilización.")

# Fase 3
add_heading("6.6. Fase 3 — Estabilización y crecimiento controlado (Semestre Ago–Dic 2027 en adelante)", level=2)
add_para("Cohorte 3 de Servicios Becarios activa.", italic=True, space_after=2)
add_para("Énfasis:", bold=True, space_after=2)
add_bullet("Operación regular del servicio a solicitud de profesores titulares del Departamento.")
add_bullet("Apertura controlada a profesores titulares de otros departamentos y a Bloques multidisciplinares.")
add_bullet("Mantenimiento y mejora incremental del portafolio.")
add_bullet("Documentación de buenas prácticas para futuras cohortes.")
add_para("Estado deseado al cierre de la fase:", bold=True, space_after=2)
add_bullet("BEER Lab consolidado como servicio recurrente del Departamento, con calendario de solicitudes estable.")
add_bullet("Capacidad instalada autosostenible: cualquier cohorte futura puede operar el laboratorio siguiendo el manual y bajo coordinación de los profesores responsables.")

# Tabla resumen
add_heading("6.7. Resumen de fases", level=2)
fases_rows = [
    ("Fase 0", "Verano 2026 (jun–ago)", "Sin becarios", "Adquisiciones, convocatoria, plantilla de ficha técnica"),
    ("Fase 1", "Semestre Ago–Dic 2026", "Cohorte 1 activa", "Portafolio base, ORSEE, mapeo, encuesta, identidad"),
    ("Receso", "Enero 2027", "Sin becarios", "Revisión, ajustes, reclutamiento cohorte 2"),
    ("Fase 2", "Semestre Feb–Jun 2027", "Cohorte 2 activa", "Pilotaje, lanzamiento, primeras sesiones, manual v1.0"),
    ("Receso", "Verano 2027 (jun–ago)", "Sin becarios", "Evaluación anual, reclutamiento cohorte 3"),
    ("Fase 3", "Semestre Ago–Dic 2027 +", "Cohorte 3 activa", "Operación estable, apertura controlada"),
]
build_table(
    ["Fase", "Calendario", "Estado del equipo", "Énfasis"],
    fases_rows,
    col_widths_cm=[2.0, 4.0, 3.0, 7.0],
)

# ====================================================================
# 7. EQUIPO DE SERVICIO BECARIO POR COHORTE SEMESTRAL
# ====================================================================
add_heading("7. Equipo de Servicio Becario por cohorte semestral", level=1)
add_para(
    "Cada cohorte semestral está compuesta por 10 a 12 estudiantes de Servicio Becario, organizados en "
    "cinco roles diferenciados. Las plazas se cubren conforme a los lineamientos institucionales del "
    "programa de Servicio Becario, con una dedicación de 80 horas durante el semestre. La cohorte está "
    "activa únicamente durante el semestre académico; en los periodos de receso (invierno y verano) los "
    "becarios no operan.",
    space_after=6,
)
add_para(
    "Tres de las plazas son ancla: la coordinación académica del portafolio, una de las plazas de "
    "desarrollo en oTree, y la plaza senior de gestión de laboratorio. Estos roles se ofrecen con "
    "expectativa explícita de continuidad multi-semestre, lo que sostiene la memoria institucional del "
    "proyecto entre cohortes. Las plazas restantes pueden rotarse por completo entre cohortes, lo que "
    "amplía el alcance formativo del programa.",
    space_after=8,
)

# 7.1 Coordinación
add_heading("7.1. Coordinación académica del portafolio — 1 plaza (ancla)", level=2)
add_para("Perfil:", bold=True, space_after=2)
add_bullet("Estudiante avanzado de Economía (5° semestre o posterior).")
add_bullet("Base sólida en microeconomía, teoría de juegos o economía conductual.")
add_bullet("Interés explícito en investigación; idealmente con experiencia previa en proyectos académicos.")
add_bullet("Disposición a permanecer en el rol durante dos o tres semestres consecutivos.")
add_para("Responsabilidades:", bold=True, space_after=2)
add_bullet("Supervisar la fidelidad teórica de cada experimento transcrito a oTree.")
add_bullet("Mantener el documento maestro del portafolio (fichas técnicas didácticas).")
add_bullet("Coordinar las reuniones del equipo y servir como enlace operativo con los profesores responsables.")
add_bullet("Validar pilotajes y firmar la lista de experimentos listos para producción.")
add_para("Producto esperado en el semestre (80 horas):", bold=True, space_after=2)
add_bullet("Fichas técnicas de seis a ocho experimentos redactadas y validadas.")
add_bullet("Contribución académica al manual de operación del laboratorio.")
add_bullet("Minutas del equipo y reporte de avance contra el cronograma.")

# 7.2 Devs
add_heading("7.2. Desarrollo en oTree — 4 plazas (1 ancla + 3 rotativas)", level=2)
add_para("Perfil:", bold=True, space_after=2)
add_bullet("Estudiante de Economía, Tecnologías de Información, Ingeniería Computacional o áreas afines.")
add_bullet("Conocimiento previo de Python o disposición verificable a aprenderlo en las primeras dos semanas.")
add_bullet("Familiaridad básica con HTML/CSS y control de versiones (git).")
add_bullet("Pensamiento estructurado y atención al detalle.")
add_para("Responsabilidades:", bold=True, space_after=2)
add_bullet("Implementar experimentos en oTree siguiendo las fichas técnicas validadas.")
add_bullet("Mantener el repositorio del laboratorio con control de versiones y documentación.")
add_bullet("Realizar testing y depuración.")
add_bullet("Resolver incidencias técnicas durante las sesiones.")
add_para("Producto esperado por desarrollador en el semestre (80 horas):", bold=True, space_after=2)
add_bullet("Dos a tres experimentos transcritos a oTree, con README en el repositorio y testing documentado.")
add_bullet("Contribuciones al mantenimiento del repositorio (refactor, librería compartida, documentación técnica).")

# 7.3 Mapeo
add_heading("7.3. Mapeo curricular y demanda — 2 plazas (rotativas)", level=2)
add_para("Perfil compartido:", bold=True, space_after=2)
add_bullet("Estudiante de Economía con habilidades de análisis de datos y comunicación escrita.")
add_bullet("Manejo de hojas de cálculo; deseable conocimiento básico de R, Python o Stata.")
add_para("Responsabilidades (analista A — mapeo curricular):", bold=True, space_after=2)
add_bullet("Revisar los syllabi de todas las UFs del Departamento y extraer temas susceptibles de experimentación.")
add_bullet("Construir la matriz UF × experimento × competencia Tec21 junto con la coordinación académica.")
add_para("Responsabilidades (analista B — encuesta y datos):", bold=True, space_after=2)
add_bullet("Diseñar y aplicar la encuesta de interés docente (Qualtrics o equivalente).")
add_bullet("Sistematizar resultados y producir el reporte ejecutivo de la encuesta.")
add_para("Producto agregado del par (160 horas combinadas):", bold=True, space_after=2)
add_bullet("Matriz curricular validada (incluye entrevistas con profesores titulares clave) y reporte de encuesta docente con tablero de respuestas, entregados antes del cierre de la Fase 1.")

# 7.4 Lab
add_heading("7.4. Gestión y asistencia de laboratorio — 3 plazas (1 ancla + 2 rotativas)", level=2)
add_para("Perfil:", bold=True, space_after=2)
add_bullet("Estudiante de Economía o áreas afines con habilidades operativas, organización y responsabilidad.")
add_bullet("Disponibilidad para trabajar en horarios de sesiones (frecuentemente vespertinos).")
add_bullet("Trato amable con participantes y manejo de imprevistos.")
add_bullet("La plaza ancla (Lab Manager Senior) sostiene la continuidad operativa entre cohortes.")
add_para("Responsabilidades del Lab Manager Senior:", bold=True, space_after=2)
add_bullet("Operar el sistema ORSEE.")
add_bullet("Reservar el laboratorio y coordinar la agenda con profesores solicitantes.")
add_bullet("Entrenar a los asistentes rotativos en el protocolo de sesión.")
add_bullet("Llevar la bitácora maestra de sesiones.")
add_para("Responsabilidades de los asistentes rotativos:", bold=True, space_after=2)
add_bullet("Preparar la sala antes de cada sesión.")
add_bullet("Ejecutar sesiones junto con el Senior: instrucciones, supervisión, soporte técnico.")
add_bullet("Procesar pagos a participantes según el protocolo institucional.")
add_para("Producto agregado por semestre (240 horas combinadas):", bold=True, space_after=2)
add_bullet("Entre doce y dieciocho sesiones preparadas y ejecutadas (pilotajes y operativas).")
add_bullet("ORSEE en operación estable con padrón actualizado al cierre del semestre.")

# 7.5 Comms
add_heading("7.5. Comunicación y difusión — 2 plazas (rotativas)", level=2)
add_para("Becario A — Diseño visual:", bold=True, space_after=2)
add_bullet("Estudiante de Comunicación, Mercadotecnia, Diseño o áreas afines.")
add_bullet("Manejo de Canva, Figma, Adobe Express o equivalente.")
add_bullet("Producto: identidad visual del laboratorio, flyer, plantillas.")
add_para("Becario B — Contenido y comunicación:", bold=True, space_after=2)
add_bullet("Estudiante de Economía, Comunicación, Letras Hispánicas o afines.")
add_bullet("Redacción clara orientada a audiencia docente.")
add_bullet("Producto: página interna informativa, materiales para la sesión informativa, comunicación con docentes y participantes vía ORSEE.")
add_para("Producto agregado del par (160 horas combinadas):", bold=True, space_after=2)
add_bullet("Identidad visual completa (logo, paleta, plantillas), flyer en versión final, página interna informativa publicada, sesión informativa al claustro coordinada y materiales asociados.")
add_bullet("Documentación de casos de uso de las primeras sesiones operativas (Fase 2 y 3).")

# 7.6 Resumen tabla
add_heading("7.6. Resumen del equipo por cohorte", level=2)
equipo_rows = [
    ("Coordinación académica del portafolio", "1", "80 h", "Avanzado de Economía", "Ancla"),
    ("Desarrollo en oTree", "4", "80 h c/u", "Economía, TI o Ingeniería con Python básico", "1 ancla + 3 rotativas"),
    ("Mapeo curricular y demanda", "2", "80 h c/u", "Economía con análisis de datos", "Rotativas"),
    ("Gestión y asistencia de laboratorio", "3", "80 h c/u", "Habilidades operativas", "1 ancla + 2 rotativas"),
    ("Comunicación y difusión", "2", "80 h c/u", "Comunicación / Diseño / Mercadotecnia", "Rotativas"),
    ("TOTAL POR COHORTE", "12", "960 h agregadas", "—", "—"),
]
build_table(
    ["Rol", "# plazas", "Dedicación", "Perfil esperado", "Continuidad"],
    equipo_rows,
    col_widths_cm=[4.5, 1.5, 2.5, 5.0, 3.0],
)

add_para(
    "Flexibilidad. La cohorte puede operar con un mínimo de 10 plazas (recortando una asistencia de "
    "laboratorio y una de comunicación) si la asignación institucional es ajustada, o expandirse hasta "
    "14 plazas si los objetivos del semestre lo requieren.",
    space_after=6,
)
add_para(
    "Continuidad entre cohortes. Las tres plazas ancla son la viga principal del proyecto: aseguran "
    "que el conocimiento acumulado en un semestre no se pierda en el siguiente. La inducción de cada "
    "nueva cohorte (primeras dos semanas del semestre) la conducen las plazas ancla, apoyadas por los "
    "profesores responsables.",
    space_after=8,
)

# ====================================================================
# 8. INDICADORES DE ÉXITO
# ====================================================================
add_heading("8. Indicadores de éxito", level=1)
ind_rows = [
    ("Número de experimentos en el portafolio", "≥ 12 al cierre de Fase 2"),
    ("Unidades de Formación con al menos un experimento asignado en el mapeo", "≥ 6 UFs prioritarias al cierre de Fase 1"),
    ("Tasa de respuesta de la encuesta de interés docente", "≥ 60 % del claustro del Departamento"),
    ("Profesores titulares que solicitan al menos una sesión durante Fase 2", "≥ 5"),
    ("Sesiones experimentales ejecutadas en Fase 2 (pilotajes más operativas)", "≥ 15"),
    ("Estudiantes-participantes únicos atendidos durante el primer año", "≥ 200"),
    ("Casos de uso documentados al cierre de Fase 2", "≥ 3 UFs distintas"),
    ("Tasa de retención de plazas ancla entre cohorte 1 y cohorte 2", "≥ 2 de 3 plazas ancla retenidas"),
    ("Componentes de know-how entregados al cierre de Fase 2", "100 % de los cinco componentes (Sección 3)"),
]
build_table(
    ["Indicador", "Meta"],
    ind_rows,
    col_widths_cm=[10.5, 5.5],
)

# ====================================================================
# 9. GOBERNANZA
# ====================================================================
add_heading("9. Gobernanza y reportes", level=1)
add_bullet("Coordinación general: profesores responsables (Castro, Maldonado), con reuniones quincenales para revisión de avance contra el cronograma e indicadores.")
add_bullet("Reunión semanal del equipo de cohorte durante el semestre, conducida por la coordinación académica del portafolio.")
add_bullet("Reporte mensual breve (una a dos páginas) a la Dirección del Departamento, con avance e indicadores.")
add_bullet("Reporte de cierre de fase al término de cada semestre, con productos entregados, indicadores cumplidos y recomendaciones para la cohorte siguiente.")
add_bullet("Repositorio único de documentación, accesible a los profesores responsables y a la cohorte activa.")
doc.add_paragraph()

# ====================================================================
# 10. RIESGOS
# ====================================================================
add_heading("10. Riesgos y mitigación", level=1)
riesgo_rows = [
    ("Baja respuesta del claustro a la encuesta de interés docente.",
     "Medio", "Alto",
     "Entrevistas uno a uno con profesores titulares clave; vincular el servicio a UFs estratégicas; mostrar pilotos exitosos como prueba de concepto en Fase 2."),
    ("Rotación completa de plazas ancla entre cohorte 1 y cohorte 2.",
     "Bajo", "Alto",
     "Reclutar plazas ancla con compromiso explícito multi-semestre; documentación exhaustiva como respaldo; reuniones de transferencia formal entre cohortes."),
    ("Curva de aprendizaje técnico de oTree en cohorte rotativa.",
     "Medio", "Medio",
     "Taller intensivo en las dos primeras semanas; pareja senior–junior en desarrollo; mentoría del desarrollador ancla a los rotativos."),
    ("Restricciones administrativas para pagar incentivos a participantes.",
     "Medio", "Alto",
     "Cerrar el acuerdo administrativo durante Fase 1, antes de cualquier sesión que requiera pago. Plan B: incentivos no monetarios para pilotajes internos."),
    ("Conflicto de uso de la sala con otras necesidades del Departamento.",
     "Bajo", "Medio",
     "Acuerdo formal de uso con la Dirección del Departamento; calendario público; reglas claras de prioridad establecidas en la Sección 4."),
    ("Priming a participantes por inexperiencia en la operación de sesiones.",
     "Medio", "Medio",
     "Capacitación específica sobre redacción de instrucciones; revisión cruzada por la coordinación académica antes de cada sesión; manual de operación con guías explícitas."),
    ("Discontinuidad entre semestres por receso de Servicios Becarios.",
     "Alto (estructural)", "Medio",
     "Plazas ancla con continuidad multi-semestre; documentación al cierre de cada cohorte; profesores responsables conducen los recesos para revisión y planeación."),
]
build_table(
    ["Riesgo", "Probabilidad", "Impacto", "Mitigación"],
    riesgo_rows,
    col_widths_cm=[4.8, 2.2, 1.8, 7.0],
)

# ====================================================================
# 11. CIERRE
# ====================================================================
add_heading("11. Cierre", level=1)
add_para(
    "La aprobación institucional del BEER Lab abre una oportunidad acotada y concreta: convertir una "
    "inversión en infraestructura en una capacidad pedagógica diferenciada del Departamento de Economía "
    "del Campus Monterrey. Esta propuesta describe la ruta operativa para lograrlo en tres semestres "
    "consecutivos, con un equipo de Servicios Becarios estructurado en torno a roles ancla que sostienen "
    "la continuidad y roles rotativos que amplían el alcance formativo del programa. Al cierre del "
    "horizonte, el Departamento contará con un servicio docente experimental consolidado, alineado al "
    "Modelo Tec21, listo para escalar hacia otros departamentos y, en el mediano plazo, hacia la "
    "vinculación con la industria vía Expedition.",
    space_after=8,
)

# ====================================================================
# ANEXO A
# ====================================================================
doc.add_page_break()
add_heading("Anexo A. Borrador de convocatoria a Servicio Becario", level=1)
add_para(
    "El siguiente texto puede adaptarse para la convocatoria oficial de cada cohorte semestral.",
    italic=True, space_after=8,
)

add_heading("Convocatoria — BEER Lab, Departamento de Economía", level=2)
add_para(
    "El Departamento de Economía del Tecnológico de Monterrey, Campus Monterrey, abre convocatoria para "
    "incorporar a 10–12 estudiantes en el equipo del Behavioral and Experimental Economics Research Lab "
    "(BEER Lab) bajo la modalidad de Servicio Becario para el próximo semestre académico. Cada plaza "
    "cubrirá 80 horas de servicio durante el semestre, distribuidas según el rol asignado.",
    space_after=8,
)
add_para("¿Qué es el BEER Lab?", bold=True, space_after=2)
add_para(
    "Un laboratorio dedicado a la docencia experimental en economía bajo el Modelo Tec21: los estudiantes "
    "participan en experimentos controlados como parte de sus Unidades de Formación, vivenciando "
    "conceptos de teoría de juegos, economía conductual y diseño de mecanismos. Es la primera "
    "infraestructura de este tipo en el Campus Monterrey y se incorpora a la oferta del Departamento "
    "como un sello distintivo del aprendizaje vivencial en economía.",
    space_after=8,
)
add_para("¿Qué harás como becario(a)?", bold=True, space_after=2)
add_para(
    "Construirás el portafolio de experimentos del laboratorio: investigarás experimentos clásicos, los "
    "programarás en oTree, los pilotearás, y luego servirás como Lab Manager o asistente durante las "
    "sesiones que los profesores titulares soliciten. Te formarás en diseño experimental, programación "
    "oTree, gestión de laboratorio y comunicación académica — competencias reconocibles en el ámbito "
    "académico y profesional.",
    space_after=8,
)
add_para("Roles disponibles (12 plazas en total por semestre):", bold=True, space_after=2)
add_bullet("Coordinación académica del portafolio — 1 plaza (con expectativa de continuidad multi-semestre)")
add_bullet("Desarrollo en oTree — 4 plazas")
add_bullet("Mapeo curricular y demanda — 2 plazas")
add_bullet("Gestión y asistencia de laboratorio — 3 plazas")
add_bullet("Comunicación y difusión — 2 plazas")
add_para("Requisitos generales:", bold=True, space_after=2)
add_bullet("Estudiante activo del Tecnológico de Monterrey, Campus Monterrey, a partir de 4° semestre (preferentemente).")
add_bullet("Promedio acumulado mínimo conforme a los lineamientos institucionales del programa de Servicio Becario.")
add_bullet("Disponibilidad de 80 horas de servicio durante el semestre académico.")
add_bullet("Interés genuino en economía, comportamiento humano o investigación aplicada.")
add_para("Cómo postular:", bold=True, space_after=2)
add_para(
    "Enviar CV breve y carta de motivación (máximo una página) a ecastrom@tec.mx y "
    "stanislao.maldonado@tec.mx, indicando el rol o roles de interés. Fecha límite: por definir cada semestre.",
    space_after=8,
)

# Guardar
doc.save(OUT)
print(f"OK: {OUT}")
